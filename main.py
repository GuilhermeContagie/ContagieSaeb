from flask import Flask, request, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import matplotlib
matplotlib.use('Agg') # Backend para execução em servidores (Render/Heroku)
import matplotlib.pyplot as plt
import numpy as np

app = Flask(__name__)

@app.route('/')
def home():
    return "API SAEB Online! Envie POST para /gerar-simulado", 200

# --- FUNÇÃO PARA RETA NUMÉRICA (MATEMÁTICA) ---
def desenhar_reta_numerica(dados):
    try:
        fig, ax = plt.subplots(figsize=(8, 2))
        min_val = dados.get('min_valor', 0)
        max_val = dados.get('max_valor', 10)
        intervalo = dados.get('intervalo_principal', 1)
        marcados = dados.get('numeros_marcados', [])
        
        ax.set_xlim(min_val - intervalo, max_val + intervalo)
        ax.set_ylim(-1, 1)
        ticks = np.arange(min_val, max_val + intervalo, intervalo)
        ax.set_xticks(ticks)
        
        labels = [str(int(t)) if t in marcados else "" for t in ticks]
        ax.set_xticklabels(labels, fontsize=12, fontweight='bold')

        destaque = dados.get('ponto_destaque')
        if destaque and destaque.get('valor') is not None:
            ax.annotate(destaque.get('rotulo', 'X'), xy=(destaque['valor'], 0), xytext=(0, 15),
                        textcoords="offset points", ha='center', color=destaque.get('cor', 'red'), 
                        fontsize=14, fontweight='bold')

        ax.spines['left'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['top'].set_visible(False)
        ax.spines['bottom'].set_position('center')
        ax.yaxis.set_visible(False)
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', bbox_inches='tight', dpi=100)
        plt.close(fig)
        img_buffer.seek(0)
        return img_buffer
    except Exception as e:
        print(f"Erro no gráfico: {e}")
        return None

# --- GERAÇÃO DO DOCUMENTO WORD ---
def criar_word_prova(dados):
    doc = Document()
    
    # Configuração de Título
    titulo = dados.get('titulo_simulado') or f"Simulado: {dados.get('materia', 'Avaliação')}"
    h = doc.add_heading(titulo, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('_' * 70)
    doc.add_paragraph('Nome: _________________________________________________ Data: ___/___/___')
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()

    # Mapeamento flexível para os dados do n8n
    itens = dados.get('itens') or dados.get('data') or dados.get('questoes') or []

    for i, item in enumerate(itens, 1):
        # 1. Identificação
        descritor = item.get('descritor_codigo', 'BNCC')
        nivel = item.get('nivel_dificuldade', '')
        doc.add_heading(f"QUESTÃO {i} ({descritor} - {nivel})", level=2)
        
        # 2. Enunciado
        enunciado = item.get('enunciado', '[Sem texto]')
        p = doc.add_paragraph(enunciado)
        p.paragraph_format.space_after = Pt(12)

        # 3. Imagem ou Gráfico
        stream_imagem = None
        
        # Caso A: Gráfico Matplotlib
        dados_python = item.get('dados_visual_python')
        if dados_python and isinstance(dados_python, dict):
            stream_imagem = desenhar_reta_numerica(dados_python)

        # Caso B: Imagem IA (Base64)
        img_b64 = item.get('imagem_base64')
        if not stream_imagem and img_b64 and len(img_b64) > 100: # Valida se é base64 real
            try:
                if "," in img_b64: img_b64 = img_b64.split(",")[1]
                img_b64 = img_b64.strip().replace("\n", "").replace("\r", "")
                
                # Corrige preenchimento (padding)
                missing_padding = len(img_b64) % 4
                if missing_padding: img_b64 += '=' * (4 - missing_padding)
                
                stream_imagem = io.BytesIO(base64.b64decode(img_b64))
            except Exception as e:
                doc.add_paragraph(f"[Erro de imagem: {str(e)}]")

        if stream_imagem:
            try:
                doc.add_picture(stream_imagem, width=Inches(3.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            except:
                doc.add_paragraph("[Falha técnica ao renderizar imagem]")

        # 4. Alternativas
        alts = item.get('alternativas', {})
        for letra in ['a', 'b', 'c', 'd']:
            if texto := alts.get(letra):
                doc.add_paragraph(f"({letra.upper()}) {texto}")
        
        doc.add_paragraph('-' * 50)

    # --- GABARITO ---
    doc.add_page_break()
    doc.add_heading('GABARITO COMENTADO', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, item in enumerate(itens, 1):
        gab = str(item.get('gabarito', '?')).upper()
        p = doc.add_paragraph()
        p.add_run(f"QUESTÃO {i}: Gabarito {gab}").bold = True
        
        just = item.get('justificativa_alternativas') or item.get('justificativa_pedagogica') or {}
        for letra, texto in just.items():
            p_j = doc.add_paragraph()
            p_j.paragraph_format.left_indent = Inches(0.3)
            p_j.add_run(f"({letra.upper()}) ").bold = True
            p_j.add_run(str(texto))
        doc.add_paragraph()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/gerar-simulado', methods=['POST'])
def gerar_simulado():
    try:
        dados = request.json
        if not dados: return {"erro": "Vazio"}, 400
            
        arquivo = criar_word_prova(dados)
        materia = str(dados.get('materia', 'Simulado')).replace(" ", "_")
        
        return send_file(
            arquivo,
            as_attachment=True,
            download_name=f"{materia}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return {"erro": str(e)}, 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
